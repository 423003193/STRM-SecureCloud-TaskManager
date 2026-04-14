import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_doc():
    doc = docx.Document()

    # Title
    title = doc.add_heading('SecureCloud Task & Resource Manager (STRM)', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Make the title purple (simulate Purple & Black theme)
    title.runs[0].font.color.rgb = RGBColor(142, 45, 226) # Accent Purple

    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph('STRM is a cross-platform Flutter application designed to work seamlessly both offline and online. It features a JARVIS-like modern, glassmorphic UI with a purple and black aesthetic.')

    doc.add_heading('2. Objectives', level=1)
    doc.add_paragraph('- Provide secure user authentication using Firebase Auth.\n- Enable robust offline task draft management using SQLite.\n- Facilitate real-time cloud data synchronization using Cloud Firestore.\n- Integrate responsive REST API data consumption.')

    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph('The app uses a clean layered architecture with Models, Services (DBHelper, AuthService, SyncService, ApiService), and Screens. It employs try-catch error handling at all layers to gracefully manage network disruptions.')

    doc.add_heading('4. Technologies Used', level=1)
    doc.add_paragraph('- Frontend: Flutter, Dart\n- Local Database: sqflite\n- Authentication: firebase_auth\n- Cloud Storage: cloud_firestore\n- API Calls: http\n- Architecture: Layered / MVVM styles')

    doc.add_heading('5. SQLite Schema', level=1)
    p = doc.add_paragraph()
    p.add_run('Table: tasks_table\n').bold = True
    p.add_run('- id: INTEGER PRIMARY KEY AUTOINCREMENT\n- title: TEXT\n- description: TEXT\n- isSynced: INTEGER DEFAULT 0')

    doc.add_heading('6. Firestore Structure', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('Collection: tasks\n').bold = True
    p2.add_run('Documents (Auto-ID):\n- title: String\n- description: String\n- createdAt: Timestamp')

    doc.add_heading('7. API Endpoints Used', level=1)
    doc.add_paragraph('GET https://jsonplaceholder.typicode.com/posts\nUsed to fetch dummy posts data to demonstrate network parsing and clean list display.')

    doc.add_heading('8. Code Implementation (Important Parts)', level=1)
    doc.add_paragraph('Core components include DBHelper.getUnsyncedTasks(), SyncService.syncTasks() for the SQLite -> Firestore flow, and comprehensive FirebaseAuthException try-catch blocks to catch invalid emails/weak passwords.')

    doc.add_heading('9. Error Handling Strategy', level=1)
    doc.add_paragraph('All async calls specifically wrap network and cloud operations in robust try-catch mechanisms. If `http.get` fails or `Firebase.initializeApp` lacks a network, the UI defaults to graceful fallback (e.g., "No Internet Connection" text widgets and SnackBar alerts).')

    doc.add_heading('10. Offline Sync Logic Explanation', level=1)
    doc.add_paragraph('1. User creates a task offline.\n2. Task saves to `sqflite` with `isSynced = 0`.\n3. Upon user clicking "Sync" (or network availability), SyncService queries SQLite for `isSynced = 0` items.\n4. Items are queued and pushed to Cloud Firestore.\n5. Pushed items are updated with `isSynced = 1` in SQLite to prevent duplication.')

    doc.add_heading('11. Screenshots (Placeholders)', level=1)
    doc.add_paragraph('[ IMAGE: Login Screen with Purple Glassmorphism ]\n[ IMAGE: Offline Task Entry Interface inside Dashboard ]\n[ IMAGE: Firestore Document verification ]\n[ IMAGE: API Data List UI ]')

    doc.add_heading('12. Output Results', level=1)
    doc.add_paragraph('The application natively handles authentication logic, queues un-synced operations when offline, completely synchronizes databases without duplication upon reconnection, and safely handles generic REST JSON payloads resulting in zero unexpected runtime faults.')

    doc.save('STRM_Project_Documentation.docx')

if __name__ == "__main__":
    create_doc()
